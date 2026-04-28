import uuid
import re
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi import Request
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from azure.data.tables import UpdateMode
from app.services.ai_evaluator import evaluate_student_answer, synthesize_final_feedback
from app.auth import get_current_user, require_course_membership, get_course_membership
from app.storage import get_scenarios_table, get_course_scenarios_table, get_labtemplates_table, utc_now_iso

router = APIRouter(prefix="/api/ai", tags=["AI"])

class AIRequest(BaseModel):
    question: str
    answer: str
    maxPoints: int
    rubric: str | None = None

@router.post("/evaluate-step")
async def ai_evaluate(data: AIRequest, current_user = Depends(get_current_user)):
    # Povoleno pro učitele, adminy i studenty (studenti volají při AUTO_SUBMIT cvičení)
    if current_user.get("global_role") not in ["teacher", "admin", "student"]:
        raise HTTPException(status_code=403, detail="Nedostatečná oprávnění")
        
    result = evaluate_student_answer(
        question=data.question,
        answer=data.answer,
        max_points=data.maxPoints,
        rubric=data.rubric
    )
    if not result:
        raise HTTPException(status_code=500, detail="AI hodnocení selhalo")

    # Pokud student nedostal plný počet bodů — přidej správnou odpověď a vysvětlení
    if result.get("correct_answer") is None and result.get("points", 0) < data.maxPoints:
        result["correct_answer"] = None
        result["explanation"] = None
        try:
            from app.services.ai_evaluator import get_correct_answer_explanation
            ca_result = get_correct_answer_explanation(
                question=data.question,
                rubric=data.rubric or ""
            )
            if ca_result:
                result["correct_answer"] = ca_result.get("correct_answer")
                result["explanation"] = ca_result.get("explanation")
        except Exception:
            pass
        
    return result

class SynthesizeRequest(BaseModel):
    feedbacks: str

@router.post("/synthesize-feedback")
async def synthesize_feedback_endpoint(data: SynthesizeRequest, current_user = Depends(get_current_user)):
    # Povoleno pro učitele, adminy i studenty (studenti volají při odevzdání AI cvičení)
    if current_user.get("global_role") not in ["teacher", "admin", "student"]:
        raise HTTPException(status_code=403, detail="Nedostatečná oprávnění")
        
    feedback = synthesize_final_feedback(data.feedbacks)
    return {"feedback": feedback}


class AdaptiveScenarioCreateRequest(BaseModel):
    title: str
    description: str = ""
    instructions: str = ""
    grading_rubric: str = ""
    required_os: str = "kali"
    time_limit: int = 60
    deadline: str | None = None
    max_attempts: int = 0
    hints: str = ""

@router.post("/courses/{course_id}/ai-scenarios")
def create_adaptive_scenario(
    course_id: str,
    payload: AdaptiveScenarioCreateRequest,
    current_user=Depends(get_current_user)
) -> JSONResponse:
    # Oprávnění: jen učitel nebo admin
    if current_user.get("global_role") not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Pouze učitel může vytvářet AI zadání.")

    membership = get_course_membership(course_id, current_user["user_id"])
    if current_user.get("global_role") != "admin":
        if not membership or membership.get("role_in_course") not in ["teacher", "assistant"]:
            raise HTTPException(status_code=403, detail="Nejste správcem tohoto kurzu.")

    # Vybereme lab template podle OS
    labtemplates = get_labtemplates_table()
    template_id = None
    for t in labtemplates.list_entities():
        if t.get("status") != "active":
            continue
        image = t.get("labImage", "")
        os_match = (
            (payload.required_os == "kali" and "adaptive-lab-kali" in image) or
            (payload.required_os == "ubuntu" and "ubuntu" in image)
        )
        if os_match:
            template_id = t.get("RowKey")
            break

    if not template_id:
        raise HTTPException(status_code=404, detail=f"Nebyl nalezen aktivní lab template pro OS '{payload.required_os}'.")

    # Vygenerujeme ID
    uid = uuid.uuid4().hex[:8]
    slug = re.sub(r"[^a-z0-9]+", "-", payload.title.lower()).strip("-")[:40]
    scenario_template_id = f"{slug}-{uid}-ai"
    course_scenario_id = f"{slug}-{uid}-cs"
    created_at = utc_now_iso()

    # 1. Uložíme scenario template
    scenarios = get_scenarios_table()
    scenarios.create_entity(entity={
        "PartitionKey": "SCENARIO_TEMPLATE",
        "RowKey": scenario_template_id,
        "scenario_template_id": scenario_template_id,
        "linked_template_id": template_id,
        "title": payload.title,
        "description": payload.description,
        "instructions": payload.instructions,
        "hints": payload.hints,
        "difficulty": "adaptive",
        "expectedOutputs": "",
        "gradingRubric": payload.grading_rubric,
        "requiredOs": payload.required_os,
        "createdBy": current_user["user_id"],
        "createdAt": created_at,
        "status": "active",
        "taskType": "adaptive",
    })

    # 2. Přiřadíme do kurzu
    course_scenarios = get_course_scenarios_table()
    course_scenarios.create_entity(entity={
        "PartitionKey": course_id,
        "RowKey": course_scenario_id,
        "courseScenarioId": course_scenario_id,
        "courseId": course_id,
        "scenarioTemplateId": scenario_template_id,
        "deadline": payload.deadline or "",
        "maxAttempts": payload.max_attempts,
        "timeLimit": payload.time_limit,
        "status": "assigned",
        "assignedBy": current_user["user_id"],
        "assignedAt": created_at,
        "taskType": "adaptive",
    })

    return JSONResponse(status_code=201, content={
        "message": "AI zadání bylo úspěšně vytvořeno.",
        "scenarioTemplateId": scenario_template_id,
        "courseScenarioId": course_scenario_id,
    })


class AdaptiveScenarioUpdateRequest(BaseModel):
    title: str
    description: str = ""
    instructions: str = ""
    grading_rubric: str = ""
    required_os: str = "kali"
    time_limit: int = 60
    deadline: str | None = None
    max_attempts: int = 0
    hints: str = ""

@router.put("/scenarios/{scenario_id}/update")
def update_adaptive_scenario(
    scenario_id: str,
    payload: AdaptiveScenarioUpdateRequest,
    current_user=Depends(get_current_user)
) -> JSONResponse:
    if current_user.get("global_role") not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Pouze učitel může upravovat AI zadání.")

    scenarios = get_scenarios_table()
    course_scenarios = get_course_scenarios_table()

    # Najdi scenario template
    try:
        entity = scenarios.get_entity(partition_key="SCENARIO_TEMPLATE", row_key=scenario_id)
    except Exception:
        raise HTTPException(status_code=404, detail="AI zadání nebylo nalezeno.")

    # Aktualizuj scenario template
    entity["title"] = payload.title
    entity["description"] = payload.description
    entity["instructions"] = payload.instructions
    entity["gradingRubric"] = payload.grading_rubric
    entity["requiredOs"] = payload.required_os
    entity["hints"] = payload.hints
    entity["updatedAt"] = utc_now_iso()
    scenarios.update_entity(entity)

    # Aktualizuj course scenario (deadline, maxAttempts, timeLimit)
    for cs in course_scenarios.list_entities():
        if cs.get("scenarioTemplateId") == scenario_id:
            cs["deadline"] = payload.deadline or ""
            cs["maxAttempts"] = payload.max_attempts
            cs["timeLimit"] = payload.time_limit
            course_scenarios.update_entity(cs)
            break

    return JSONResponse(status_code=200, content={"message": "AI zadání bylo úspěšně aktualizováno."})


# ─── Materiály AI scénáře ────────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg", "docx", "pptx", "mp4", "txt", "md"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

def get_ai_blob_service():
    import os
    from azure.storage.blob import BlobServiceClient
    conn = os.getenv("BLOB_CONNECTION_STRING")
    if not conn:
        raise RuntimeError("Chybí BLOB_CONNECTION_STRING v .env")
    return BlobServiceClient.from_connection_string(conn)

def get_ai_materials_table():
    import os
    from app.storage import ensure_table
    return ensure_table(os.getenv("MATERIALS_TABLE_NAME", "coursematerials"))

def get_blob_container_name():
    import os
    return os.getenv("BLOB_CONTAINER_MATERIALS", "course-materials")


@router.post("/scenarios/{scenario_id}/materials")
async def upload_ai_scenario_material(
    scenario_id: str,
    file: UploadFile = File(...),
    current_user=Depends(get_current_user)
):
    if current_user.get("global_role") not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Nedostatečná oprávnění.")

    original_name = file.filename or "soubor"
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Nepodporovaný typ souboru. Povoleno: {', '.join(ALLOWED_EXTENSIONS)}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Soubor je příliš velký. Maximum je 20 MB.")

    file_id = uuid.uuid4().hex[:12]
    blob_name = f"ai-scenarios/{scenario_id}/{file_id}.{ext}"

    blob_service = get_ai_blob_service()
    container_name = get_blob_container_name()
    blob_client = blob_service.get_blob_client(container=container_name, blob=blob_name)
    blob_client.upload_blob(content, overwrite=True)

    table = get_ai_materials_table()
    uploaded_at = utc_now_iso()
    table.create_entity(entity={
        "PartitionKey": f"ai-scenario-{scenario_id}",
        "RowKey": file_id,
        "fileId": file_id,
        "scenarioId": scenario_id,
        "originalName": original_name,
        "blobName": blob_name,
        "extension": ext,
        "sizeBytes": len(content),
        "uploadedBy": current_user["user_id"],
        "uploadedAt": uploaded_at,
    })

    return JSONResponse(status_code=201, content={
        "message": "Soubor úspěšně nahrán.",
        "fileId": file_id,
        "originalName": original_name,
        "sizeBytes": len(content),
        "uploadedAt": uploaded_at,
    })


@router.get("/scenarios/{scenario_id}/materials")
def list_ai_scenario_materials(
    scenario_id: str,
    current_user=Depends(get_current_user)
):
    table = get_ai_materials_table()
    results = []
    for entity in table.query_entities(query_filter=f"PartitionKey eq 'ai-scenario-{scenario_id}'"):
        results.append({
            "fileId": entity.get("fileId"),
            "originalName": entity.get("originalName"),
            "extension": entity.get("extension"),
            "sizeBytes": entity.get("sizeBytes"),
            "uploadedBy": entity.get("uploadedBy"),
            "uploadedAt": entity.get("uploadedAt"),
        })
    results.sort(key=lambda x: x.get("uploadedAt", ""), reverse=True)
    return results


@router.delete("/scenarios/{scenario_id}/materials/{file_id}")
def delete_ai_scenario_material(
    scenario_id: str,
    file_id: str,
    current_user=Depends(get_current_user)
):
    if current_user.get("global_role") not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Nedostatečná oprávnění.")

    table = get_ai_materials_table()
    try:
        from azure.core.exceptions import ResourceNotFoundError
        entity = table.get_entity(partition_key=f"ai-scenario-{scenario_id}", row_key=file_id)
    except Exception:
        raise HTTPException(status_code=404, detail="Soubor nebyl nalezen.")

    blob_name = entity.get("blobName")
    blob_service = get_ai_blob_service()
    container_name = get_blob_container_name()
    try:
        blob_service.get_blob_client(container=container_name, blob=blob_name).delete_blob()
    except Exception:
        pass

    table.delete_entity(partition_key=f"ai-scenario-{scenario_id}", row_key=file_id)
    return {"message": f"Soubor {file_id} byl smazán."}


@router.get("/scenarios/{scenario_id}/materials/content")
async def get_ai_scenario_materials_content(
    scenario_id: str,
    current_user=Depends(get_current_user)
):
    """Vrátí textový obsah všech materiálů scénáře pro použití v AI promptu."""
    import os
    import io

    table = get_ai_materials_table()
    blob_service = get_ai_blob_service()
    container_name = get_blob_container_name()

    results = []
    for entity in table.query_entities(query_filter=f"PartitionKey eq 'ai-scenario-{scenario_id}'"):
        blob_name = entity.get("blobName", "")
        ext = entity.get("extension", "").lower()
        original_name = entity.get("originalName", "soubor")

        # Pouze textové formáty
        if ext not in {"pdf", "txt", "md", "docx"}:
            continue

        try:
            blob_client = blob_service.get_blob_client(container=container_name, blob=blob_name)
            data = blob_client.download_blob().readall()

            text = ""
            if ext in {"txt", "md"}:
                text = data.decode("utf-8", errors="ignore")
            elif ext == "pdf":
                try:
                    import pypdf
                    reader = pypdf.PdfReader(io.BytesIO(data))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                except Exception:
                    text = ""
            elif ext == "docx":
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(data))
                    text = "\n".join(p.text for p in doc.paragraphs)
                except Exception:
                    text = ""

            if text.strip():
                results.append({
                    "fileName": original_name,
                    "content": text[:15000]  # max 15k znaků na soubor
                })
        except Exception:
            continue

    return {"materials": results}


class AIChatRequest(BaseModel):
    system: str
    message: str

@router.post("/chat")
async def ai_chat(data: AIChatRequest, current_user=Depends(get_current_user)):
    """Obecný AI chat endpoint pro studentský portál (generování intro, podúkolů, hodnocení)."""
    from app.services.ai_evaluator import get_ai_client
    import json

    client = get_ai_client()
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": data.system},
                {"role": "user",   "content": data.message},
            ]
        )
        return {"response": response.choices[0].message.content.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI chyba: {str(e)}")


class AIEvaluateSubtaskRequest(BaseModel):
    system: str
    message: str
    max_points: int = Field(alias="maxPoints")

@router.post("/evaluate-subtask")
async def ai_evaluate_subtask(data: AIEvaluateSubtaskRequest, current_user=Depends(get_current_user)):
    """Hodnocení odpovědi studenta na AI podúkol — vrací JSON {points, feedback}."""
    from app.services.ai_evaluator import get_ai_client
    import json

    client = get_ai_client()
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": data.system},
                {"role": "user",   "content": data.message},
            ],
            response_format={"type": "json_object"}
        )
        raw = response.choices[0].message.content.strip()
        parsed = json.loads(raw)
        points = max(0, min(int(data.max_points), int(round(float(parsed.get("points", 0))))))
        return {
            "points": points,
            "feedback": str(parsed.get("feedback", "")).strip(),
            "correct_answer": parsed.get("correct_answer") or None,
            "explanation": parsed.get("explanation") or None,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI chyba: {str(e)}")


# ── Init script generator ────────────────────────────────────────────────────

class InitScriptGenerateRequest(BaseModel):
    base_image: str = "kali"
    goal: str = ""
    files: list[str] = []
    file_paths: dict = {}   # { "malware.exe": "/root/lab/malware.exe", ... }

@router.post("/generate-init-script")
async def generate_init_script(
    data: InitScriptGenerateRequest,
    current_user=Depends(get_current_user),
) -> JSONResponse:
    if current_user.get("global_role") not in ["teacher", "admin"]:
        raise HTTPException(status_code=403, detail="Nedostatečná oprávnění")

    from app.services.ai_evaluator import get_ai_client
    client = get_ai_client()

    base_label = "Kali Linux (GUI)" if data.base_image == "kali" else "Ubuntu (CLI)"

    files_block = ""
    if data.files:
        lines = []
        for f in data.files:
            target = data.file_paths.get(f, "").strip()
            if target:
                lines.append(f"  - {f}  →  umístit do: {target}")
            else:
                lines.append(f"  - {f}  (cílová cesta neurčena)")
        files_block = "Nahrané soubory v blob storage:\n" + "\n".join(lines)
    else:
        files_block = "Žádné další soubory nebyly nahrány."

    system_prompt = (
        "Jsi expert na tvorbu bash init scriptů pro kybernetické laboratorní prostředí na univerzitě.\n\n"
        "Platforma funguje takto:\n"
        "- Lab kontejner se spouští z Docker image (Kali Linux nebo Ubuntu)\n"
        "- Při každém startu kontejneru se stáhne a spustí init.sh pomocí env proměnné LAB_INIT_BLOB_URL\n"
        "- Ostatní soubory jsou uloženy ve stejné složce Azure Blob Storage jako init.sh\n"
        "- Pro stažení ostatních souborů ze stejné složky použij tento pattern:\n"
        "    FILE_URL=\"${LAB_INIT_BLOB_URL/init.sh/nazev_souboru}\"\n"
        "    curl -fsSL -o /cil/souboru \"$FILE_URL\"\n"
        "- Kontejner MÁ přístup k internetu (curl, apt-get, wget fungují)\n"
        "- Skript běží jako root\n\n"
        "Pravidla:\n"
        "1. Začni #!/bin/bash a nastav 'set -e' pro okamžité selhání při chybě\n"
        "2. Každý logický blok odděluj komentářem\n"
        "3. Vytvoř cílové adresáře před kopírováním (mkdir -p)\n"
        "4. Nastav práva chmod dle účelu (spustitelné soubory 755, citlivé soubory 600)\n"
        "5. Skript musí být idempotentní (bezpečné spustit vícekrát)\n"
        "6. VRAŤ POUZE obsah bash skriptu — žádný markdown, žádné kód bloky, žádné vysvětlivky vně komentářů"
    )

    user_prompt = (
        f"Base image: {base_label}\n"
        f"{files_block}\n\n"
        f"Popis labu / co má student v labu dělat:\n{data.goal or 'Obecné kybernetické lab prostředí.'}\n\n"
        "Vygeneruj kompletní init.sh pro tento lab."
    )

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.25,
            max_tokens=1800,
        )
        script = response.choices[0].message.content.strip()
        if script.startswith("```"):
            script = re.sub(r"^```[a-z]*\n?", "", script)
            script = re.sub(r"\n?```$", "", script.strip())
        return JSONResponse(content={"init_script": script})
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"AI chyba: {exc}")